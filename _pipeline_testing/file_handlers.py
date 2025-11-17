"""File I/O and text extraction utilities."""

import hashlib
import zipfile
import re
from pathlib import Path
from typing import Optional
import PyPDF2
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub
    EPUBLIB_AVAILABLE = True
except ImportError:
    EPUBLIB_AVAILABLE = False


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text content from PDF file."""
    text_parts = []
    try:
        with open(pdf_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_txt(txt_path: Path) -> str:
    """Extract text content from TXT file."""
    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(txt_path, "r", encoding="latin-1") as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {e}")


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text content from DOCX file."""
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
    
    try:
        doc = Document(docx_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_epub(epub_path: Path) -> str:
    """Extract text content from EPUB file."""
    if EPUBLIB_AVAILABLE:
        # Use ebooklib for better EPUB parsing
        try:
            book = epub.read_epub(str(epub_path))
            text_parts = []
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Get the content as HTML
                    content = item.get_content().decode('utf-8')
                    # Simple HTML tag removal - extract text between tags
                    # Remove script and style tags and their content
                    content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
                    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
                    # Remove HTML tags
                    content = re.sub(r'<[^>]+>', '', content)
                    # Decode HTML entities
                    import html
                    content = html.unescape(content)
                    # Clean up whitespace
                    content = re.sub(r'\s+', ' ', content).strip()
                    if content:
                        text_parts.append(content)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from EPUB using ebooklib: {e}")
    else:
        # Fallback: Use zipfile to extract HTML files and parse them
        try:
            text_parts = []
            with zipfile.ZipFile(epub_path, 'r') as zip_ref:
                # EPUB structure: find HTML/XHTML files in OEBPS or similar folders
                html_files = [f for f in zip_ref.namelist() if f.endswith(('.html', '.xhtml', '.htm')) and 'META-INF' not in f and 'mimetype' not in f]
                
                for html_file in html_files:
                    try:
                        content = zip_ref.read(html_file).decode('utf-8')
                        # Remove script and style tags
                        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
                        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
                        # Remove HTML tags
                        content = re.sub(r'<[^>]+>', '', content)
                        # Decode HTML entities
                        import html
                        content = html.unescape(content)
                        # Clean up whitespace
                        content = re.sub(r'\s+', ' ', content).strip()
                        if content:
                            text_parts.append(content)
                    except Exception:
                        continue  # Skip files that can't be decoded
            
            if not text_parts:
                raise ValueError("No readable HTML content found in EPUB file")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from EPUB: {e}. Install ebooklib for better support: pip install ebooklib")


def extract_text_from_file(file_path: Path) -> str:
    """Extract text from file based on extension."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".epub":
        return extract_text_from_epub(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def calculate_file_hash(file_path: Path, algorithm: str = "sha256") -> str:
    """Calculate hash of a file."""
    hash_obj = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    except Exception as e:
        raise ValueError(f"Failed to calculate hash: {e}")


def calculate_json_hash(data: dict, algorithm: str = "sha256") -> str:
    """Calculate hash of a JSON-serializable object."""
    import json
    # Serialize to JSON with sorted keys for deterministic hashing
    json_str = json.dumps(data, sort_keys=True, ensure_ascii=False)
    hash_obj = hashlib.sha256()
    hash_obj.update(json_str.encode('utf-8'))
    return hash_obj.hexdigest()


